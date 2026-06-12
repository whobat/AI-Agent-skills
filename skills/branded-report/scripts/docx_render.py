# -*- coding: utf-8 -*-
"""
Render a Markdown report to a themed .docx using python-docx + beautifulsoup4.

Markdown is converted to HTML (via `markdown`) and the HTML is walked to emit Word
content: a branded cover, headings (h1 white-on-primary band, h2 primary, h3 secondary),
paragraphs with bold/italic/inline-code, bullet/numbered lists, tables (shaded header +
banded rows), fenced code blocks, blockquotes, and rules. Colors and fonts come from the
same theme dict the HTML/PDF path uses, so all three formats match.
"""
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Mm


def _hex(color):
    return RGBColor.from_string(color.lstrip("#").upper())


def _font_name(font_value):
    # theme fonts may be CSS stacks like "Georgia, serif" -> take the first family
    return font_value.split(",")[0].strip().strip("'\"")


def _shade(element, fill_hex):
    """Apply background shading to a paragraph (_p) or table cell (_tc) element."""
    pr = element.get_or_add_pPr() if element.tag.endswith("}p") else element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#").upper())
    pr.append(shd)


def _bottom_border(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color_hex.lstrip("#").upper())
    pbdr.append(bottom)
    pPr.append(pbdr)


class DocxBuilder:
    def __init__(self, theme):
        self.c = theme["colors"]
        self.f = theme["fonts"]
        self.heading_font = _font_name(self.f["heading"])
        self.body_font = _font_name(self.f["body"])
        self.mono_font = _font_name(self.f["mono"])
        self.doc = Document()
        # Base body font on Normal style.
        normal = self.doc.styles["Normal"]
        normal.font.name = self.body_font
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = _hex(self.c["text"])

    # ---- inline runs -----------------------------------------------------
    def _add_inline(self, paragraph, node, bold=False, italic=False, mono=False):
        if isinstance(node, NavigableString):
            text = str(node)
            if not text:
                return
            run = paragraph.add_run(text)
            run.bold = bold or None
            run.italic = italic or None
            if mono:
                run.font.name = self.mono_font
                run.font.size = Pt(9)
                run.font.color.rgb = _hex(self.c["muted"])
            elif bold:
                run.font.color.rgb = _hex(self.c["muted"])
            return
        if isinstance(node, Tag):
            b = bold or node.name in ("strong", "b")
            i = italic or node.name in ("em", "i")
            m = mono or node.name == "code"
            for child in node.children:
                self._add_inline(paragraph, child, b, i, m)

    def _add_inline_children(self, paragraph, tag):
        for child in tag.children:
            self._add_inline(paragraph, child)

    # ---- blocks ----------------------------------------------------------
    def heading1(self, tag):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        _shade(p._p, self.c["primary"])
        run = p.add_run(tag.get_text())
        run.bold = True
        run.font.name = self.heading_font
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def heading2(self, tag):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(tag.get_text())
        run.bold = True
        run.font.name = self.heading_font
        run.font.size = Pt(13)
        run.font.color.rgb = _hex(self.c["primary"])
        _bottom_border(p, self.c["border"])

    def heading3(self, tag):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(tag.get_text())
        run.bold = True
        run.font.name = self.heading_font
        run.font.size = Pt(11)
        run.font.color.rgb = _hex(self.c["secondary"])

    def paragraph(self, tag):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        self._add_inline_children(p, tag)

    def listing(self, tag, ordered, level=0):
        style = "List Number" if ordered else "List Bullet"
        for li in tag.find_all("li", recursive=False):
            p = self.doc.add_paragraph(style=style)
            if level:
                p.paragraph_format.left_indent = Mm(6 * (level + 1))
            # inline content (everything that's not a nested list)
            for child in li.children:
                if isinstance(child, Tag) and child.name in ("ul", "ol"):
                    continue
                self._add_inline(p, child)
            for sub in li.find_all(["ul", "ol"], recursive=False):
                self.listing(sub, sub.name == "ol", level + 1)

    def code_block(self, tag):
        text = tag.get_text().rstrip("\n")
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        _shade(p._p, self.c["light"])
        for j, line in enumerate(text.split("\n")):
            if j:
                p.add_run().add_break()
            run = p.add_run(line)
            run.font.name = self.mono_font
            run.font.size = Pt(8.5)
            run.font.color.rgb = _hex(self.c["text"])

    def blockquote(self, tag):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Mm(6)
        p.paragraph_format.space_after = Pt(5)
        _shade(p._p, self.c["light"])
        run = p.add_run(tag.get_text().strip())
        run.italic = True
        run.font.color.rgb = _hex(self.c["muted"])

    def table(self, tag):
        rows = tag.find_all("tr")
        if not rows:
            return
        ncols = max(len(r.find_all(["td", "th"])) for r in rows)
        t = self.doc.add_table(rows=0, cols=ncols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            t.style = "Table Grid"
        except KeyError:
            pass
        for ri, tr in enumerate(rows):
            cells = tr.find_all(["td", "th"])
            is_header = tr.find("th") is not None
            row = t.add_row()
            for ci in range(ncols):
                cell = row.cells[ci]
                cell.paragraphs[0].text = ""
                src = cells[ci] if ci < len(cells) else None
                para = cell.paragraphs[0]
                if src is not None:
                    self._add_inline_children(para, src)
                for run in para.runs:
                    run.font.size = Pt(9)
                if is_header:
                    _shade(cell._tc, self.c["secondary"])
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif ri % 2 == 0:
                    _shade(cell._tc, self.c["light"])

    def horizontal_rule(self):
        p = self.doc.add_paragraph()
        _bottom_border(p, self.c["border"])

    # ---- cover -----------------------------------------------------------
    def cover(self, theme, theme_path, title, subtitle, meta_lines):
        # Word cannot embed SVG: use logo_raster (PNG/JPG) if given, else logo when it is a raster.
        def _resolve(val):
            if not val:
                return None
            p = Path(str(val).replace("\\", "/"))
            if not p.is_absolute() and theme_path:
                p = (Path(theme_path).parent / str(val).replace("\\", "/")).resolve()
            return p if p.exists() else None
        raster = _resolve(theme.get("logo_raster"))
        if raster is None:
            cand = _resolve(theme.get("logo"))
            if cand is not None and cand.suffix.lower() in (".png", ".jpg", ".jpeg"):
                raster = cand
        if raster is not None:
            try:
                self.doc.add_picture(str(raster), width=Mm(55))
            except Exception:  # noqa: BLE001
                pass
        org = theme.get("organization")
        if org:
            op = self.doc.add_paragraph()
            r = op.add_run(org)
            r.bold = True
            r.font.name = self.heading_font
            r.font.size = Pt(12)
            r.font.color.rgb = _hex(self.c["secondary"])
        tp = self.doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(10)
        tr = tp.add_run(title)
        tr.bold = True
        tr.font.name = self.heading_font
        tr.font.size = Pt(26)
        tr.font.color.rgb = _hex(self.c["primary"])
        if subtitle:
            sp = self.doc.add_paragraph()
            sr = sp.add_run(subtitle)
            sr.italic = True
            sr.font.name = self.heading_font
            sr.font.size = Pt(14)
            sr.font.color.rgb = _hex(self.c["secondary"])
        if meta_lines:
            self.doc.add_paragraph()
            for line in meta_lines:
                # meta lines may contain simple <b>..</b> from the inline-markdown helper
                frag = BeautifulSoup(line, "html.parser")
                mp = self.doc.add_paragraph()
                mp.paragraph_format.space_after = Pt(2)
                self._add_inline_children(mp, frag)
                for run in mp.runs:
                    run.font.color.rgb = _hex(self.c["muted"])
        self.doc.add_page_break()

    # ---- driver ----------------------------------------------------------
    def walk(self, soup):
        handlers = {
            "h1": self.heading1, "h2": self.heading2, "h3": self.heading3,
            "h4": self.heading3, "h5": self.heading3, "h6": self.heading3,
            "p": self.paragraph, "table": self.table, "blockquote": self.blockquote,
        }
        for el in soup.find_all(True, recursive=False) if soup.name else soup.children:
            self._dispatch(el, handlers)
        # top-level body children
    def _dispatch(self, el, handlers):
        if isinstance(el, NavigableString):
            return
        name = el.name
        if name in handlers:
            handlers[name](el)
        elif name == "ul":
            self.listing(el, ordered=False)
        elif name == "ol":
            self.listing(el, ordered=True)
        elif name == "pre":
            self.code_block(el)
        elif name == "hr":
            self.horizontal_rule()


def render_docx(md_text, theme, theme_path, title, subtitle, meta_lines, out_path):
    html = markdown.Markdown(extensions=["tables", "fenced_code", "sane_lists"]).convert(md_text)
    soup = BeautifulSoup(html, "html.parser")
    b = DocxBuilder(theme)
    if title:
        b.cover(theme, theme_path, title, subtitle, meta_lines)
    handlers = {
        "h1": b.heading1, "h2": b.heading2, "h3": b.heading3,
        "h4": b.heading3, "h5": b.heading3, "h6": b.heading3,
        "p": b.paragraph, "table": b.table, "blockquote": b.blockquote,
    }
    for el in soup.children:
        b._dispatch(el, handlers)
    b.doc.save(out_path)
    return out_path
