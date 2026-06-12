#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tests for build_report.py / docx_render.py / extract_theme.py.

  python -m unittest test_branded_report -v

Hermetic: builds a tiny Markdown + theme in a temp dir and asserts on the outputs.
HTML and DOCX are always exercised. PDF is only checked when a headless browser is
found (skipped otherwise). extract_theme is tested against a hand-built minimal pptx.
"""
import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
import build_report as br  # noqa: E402
import extract_theme as et  # noqa: E402

MD = """# Heading One

## Heading Two

Body with **bold**, *italic* and `code`.

- bullet a
- bullet b

| Col A | Col B |
|-------|-------|
| x | 1 |
| y | 2 |

```sql
SELECT 1;
```

> A quote.
"""

THEME = {
    "organization": "Acme Corp",
    "colors": {"primary": "#AA0000", "secondary": "#006600", "accent": "#88CC88",
               "light": "#F0F0F0", "border": "#CCCCCC", "muted": "#555555", "text": "#222222"},
    "fonts": {"heading": "Georgia, serif", "body": "Georgia, serif", "mono": "Consolas, monospace"},
    "logo": None,
}


class TestHtml(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.md = Path(self.tmp) / "r.md"
        self.md.write_text(MD, encoding="utf-8")
        self.theme = Path(self.tmp) / "theme.json"
        self.theme.write_text(json.dumps(THEME), encoding="utf-8")

    def test_html_contains_theme_and_content(self):
        br.main(["--input", str(self.md), "--theme", str(self.theme),
                 "--formats", "html", "--output-dir", self.tmp, "--name", "out",
                 "--title", "My Report", "--subtitle", "Sub"])
        html = (Path(self.tmp) / "out.html").read_text(encoding="utf-8")
        self.assertIn("#AA0000", html)            # primary colour wired into CSS
        self.assertIn("My Report", html)          # cover title
        self.assertIn("<table>", html)            # table rendered
        self.assertIn("<h1", html)                # heading from markdown
        self.assertIn("Acme Corp", html)          # organization

    def test_default_theme_when_no_theme_given(self):
        br.main(["--input", str(self.md), "--formats", "html",
                 "--output-dir", self.tmp, "--name", "d"])
        html = (Path(self.tmp) / "d.html").read_text(encoding="utf-8")
        self.assertIn(br.DEFAULT_THEME["colors"]["primary"], html)


class TestDocx(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.md = Path(self.tmp) / "r.md"
        self.md.write_text(MD, encoding="utf-8")
        self.theme = Path(self.tmp) / "theme.json"
        self.theme.write_text(json.dumps(THEME), encoding="utf-8")

    def test_docx_produced_with_structure(self):
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed")
        br.main(["--input", str(self.md), "--theme", str(self.theme),
                 "--formats", "docx", "--output-dir", self.tmp, "--name", "out",
                 "--title", "My Report", "--subtitle", "Sub"])
        path = Path(self.tmp) / "out.docx"
        self.assertTrue(path.exists() and path.stat().st_size > 0)
        import docx
        d = docx.Document(str(path))
        self.assertEqual(len(d.tables), 1)
        self.assertEqual(len(d.tables[0].rows), 3)          # header + 2 data rows
        texts = "\n".join(p.text for p in d.paragraphs)
        self.assertIn("Heading One", texts)
        self.assertIn("My Report", texts)                   # cover title present
        # the H1 run is white-on-primary
        h1 = [p for p in d.paragraphs if p.text == "Heading One"][0]
        self.assertEqual(str(h1.runs[0].font.color.rgb), "FFFFFF")


class TestLogoResolution(unittest.TestCase):
    PNG_1x1 = bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c489"
        "0000000d49444154789c6360000002000001e221bc330000000049454e44ae426082")

    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())
        (self.tmp / "assets").mkdir()
        (self.tmp / "assets" / "logo.svg").write_text(
            '<svg xmlns="http://www.w3.org/2000/svg"/>', encoding="utf-8")
        (self.tmp / "assets" / "logo.png").write_bytes(self.PNG_1x1)
        self.theme_path = str(self.tmp / "theme.json")

    def test_relative_asset_resolves_against_theme_folder(self):
        # forward-slash relative path
        theme = {"logo": "assets/logo.png"}
        Path(self.theme_path).write_text(json.dumps(theme), encoding="utf-8")
        data_uri, raster = br.resolve_logo(theme, self.theme_path)
        self.assertTrue(data_uri.startswith("data:image/png;base64,"))
        self.assertTrue(raster.endswith("logo.png"))

    def test_backslash_path_is_normalized(self):
        theme = {"logo": ".\\assets\\logo.png"}
        _, raster = br.resolve_logo(theme, self.theme_path)
        self.assertIsNotNone(raster)
        self.assertTrue(raster.endswith("logo.png"))

    def test_svg_logo_with_png_logo_raster_for_docx(self):
        # HTML/PDF get the crisp SVG; DOCX gets the PNG raster.
        theme = {"logo": "assets/logo.svg", "logo_raster": "assets/logo.png"}
        data_uri, raster = br.resolve_logo(theme, self.theme_path)
        self.assertTrue(data_uri.startswith("data:image/svg+xml;base64,"))
        self.assertTrue(raster.endswith("logo.png"))

    def test_svg_only_gives_no_docx_raster(self):
        theme = {"logo": "assets/logo.svg"}
        data_uri, raster = br.resolve_logo(theme, self.theme_path)
        self.assertTrue(data_uri.startswith("data:image/svg+xml"))
        self.assertIsNone(raster)


class TestExtractTheme(unittest.TestCase):
    def _make_pptx(self, path):
        theme_xml = (
            '<?xml version="1.0"?>'
            '<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="T">'
            '<a:themeElements><a:clrScheme name="c">'
            '<a:dk1><a:srgbClr val="000000"/></a:dk1><a:lt1><a:srgbClr val="FFFFFF"/></a:lt1>'
            '<a:dk2><a:srgbClr val="6B5849"/></a:dk2><a:lt2><a:srgbClr val="FFFFFF"/></a:lt2>'
            '<a:accent1><a:srgbClr val="AD1922"/></a:accent1>'
            '<a:accent2><a:srgbClr val="50634D"/></a:accent2>'
            '<a:accent3><a:srgbClr val="92C56E"/></a:accent3>'
            '<a:accent4><a:srgbClr val="CFAD87"/></a:accent4>'
            '<a:accent5><a:srgbClr val="E7D6C3"/></a:accent5>'
            '<a:accent6><a:srgbClr val="F1E5D5"/></a:accent6>'
            '</a:clrScheme>'
            '<a:fontScheme name="f">'
            '<a:majorFont><a:latin typeface="Georgia"/></a:majorFont>'
            '<a:minorFont><a:latin typeface="Georgia"/></a:minorFont>'
            '</a:fontScheme></a:themeElements></a:theme>'
        )
        with zipfile.ZipFile(path, "w") as z:
            z.writestr("ppt/theme/theme1.xml", theme_xml)
            z.writestr("ppt/media/image1.svg", '<svg xmlns="http://www.w3.org/2000/svg"/>')

    def test_extracts_colors_and_fonts(self):
        tmp = tempfile.mkdtemp()
        pptx = Path(tmp) / "Brand_Template.pptx"
        self._make_pptx(pptx)
        theme = et.extract(str(pptx), str(Path(tmp) / "assets"))
        self.assertEqual(theme["colors"]["primary"], "#AD1922")    # accent1
        self.assertEqual(theme["colors"]["secondary"], "#50634D")  # accent2
        self.assertEqual(theme["colors"]["muted"], "#6B5849")      # dk2
        self.assertIn("Georgia", theme["fonts"]["heading"])
        self.assertEqual(theme["organization"], "Brand Template")  # stem, underscores->spaces
        self.assertTrue(theme["logo"].endswith("image1.svg"))      # logo candidate copied


if __name__ == "__main__":
    unittest.main(verbosity=2)
